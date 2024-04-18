# Eduardo Enzo Pedro de Melo - 16/04/2024
from os import system
import pandas as pd
from docx import Document
from docx.shared import Pt

system('cls')

df = pd.read_excel('D:\pythonExcel\pythonExcel\projetoReal\Tabela Projeto.xlsx')
#print(df)


contador = 1

for index, rows in df.iterrows():
    
    document = Document()

    document.settings
    document.add_heading('Curriculo', 0)
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(14)

    nome, sobrenome, dataNascimento, graduacao, universidade, anoGraduacao, areaAtuacao = rows # Unpacking

    p = document.add_paragraph(f'   {nome} {sobrenome}, nascido em {dataNascimento}, cursou {graduacao} em {universidade}. Concluiu o curso em {anoGraduacao}. Atualmente, {nome} atua na área de {areaAtuacao}')    
        
    document.save(f'D:\pythonExcel\pythonExcel\projetoReal\curriculos\currículo - {contador} ({nome} {sobrenome}) .docx')
    contador += 1

    #print(f'{nome} {sobrenome}, nascido em {dataDeNascimento}, cursou {graduacao} na universidade {universidade}. Concluiu o curso no ano de {anoGraduacao}. Altualmente, {nome} atua na área de {areaAtuacao}\n')
